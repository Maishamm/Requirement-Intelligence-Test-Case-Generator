from __future__ import annotations

from pathlib import Path
from typing import Optional


def load_requirement_text(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Requirement file not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix in [".txt", ".md"]:
        return _load_text_file(path)
    if suffix == ".docx":
        return _load_docx_file(path)
    if suffix == ".pdf":
        return _load_pdf_file(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def _load_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_docx_file(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Please install python-docx: pip install python-docx") from None

    doc = Document(path)  # type: ignore
    parts: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # Tables often contain mandatory fields, validations, acceptance criteria, and status rules.
    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"[DOCX Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _load_pdf_file(path: Path) -> str:
    text = _try_pypdf(path)
    if not text or len(text.strip()) < 50:
        text = _try_pdfplumber(path)
    return (text or "").strip()


def _try_pypdf(path: Path) -> Optional[str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        print("[WARN] pypdf is not installed. Will try fallback.")
        return None

    try:
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text.strip())
        return "\n".join(texts)
    except Exception as e:
        print(f"[WARN] pypdf failed: {e}")
        return None


def _try_pdfplumber(path: Path) -> Optional[str]:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Please install pdfplumber: pip install pdfplumber") from None

    try:
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text.strip())
        return "\n".join(texts)
    except Exception as e:
        print(f"[WARN] pdfplumber failed: {e}")
        return None
